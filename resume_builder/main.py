"""
简历智能生成器 — 完整流水线
PDF/TXT简历 → 千问分析+STAR优化 → 本地模板 + AI填入 → 输出专业 .docx
"""
import json, os, sys
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from qwen_client import call_qwen

# 项目目录结构
BASE_DIR = Path(__file__).parent
RESUMES_DIR = BASE_DIR / "resumes"          # 放原始简历（.pdf / .txt）
TEMPLATES_DIR = BASE_DIR / "templates"      # 放本地模板（.docx）
OUTPUT_DIR = BASE_DIR / "output"            # 输出生成后的简历
for d in [RESUMES_DIR, TEMPLATES_DIR, OUTPUT_DIR]:
    d.mkdir(exist_ok=True)

# ============ 工具函数 ============

def clean_json(text: str) -> str:
    """清洗模型返回的JSON：去markdown fence、去前后废话"""
    text = text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[-1].rsplit("\n```", 1)[0] if "\n```" in text else text[3:-3]
    return text.strip()


def safe_json_parse(text: str, default=None):
    """安全解析JSON，支持从文本中提取"""
    text = clean_json(text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # 尝试找第一个 { 到最后一个 }
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end > start:
            try:
                return json.loads(text[start:end + 1])
            except json.JSONDecodeError:
                pass
        return default if default is not None else {}


# ============ 第1步：分析 + 优化简历 ============

ANALYZE_PROMPT = """你是一位资深HR和职业简历优化专家，有10年以上招聘经验。你知道HR平均6秒扫一份简历，所以每个字都必须有说服力。

请分析以下应届生简历，输出弱点诊断和优化后的完整简历。

=== 核心原则 ===
1. 量化优先：所有成果都要有数字。原文有的精确保留，原文没有的根据上下文合理推估
2. 动词开头：使用"主导/设计/实现/优化/搭建/重构/攻克/推动"等强动词，禁用"参与/负责"等弱描述
3. 关键词匹配：预判JD中的ATS关键词（技术栈、工具、方法论）并自然融入
4. 结果导向：每条经历的落点必须是可量化的业务价值
5. 精简原则：每条要点控制在30字以内
6. ⚠️ 真实性第一：公司名、职位、时间等事实信息必须原文照搬，绝不编造。AI实习就是AI实习，不要变成阿里云实习。

=== 个人概述（summary）撰写规则 ===
不要写"我热爱学习、吃苦耐劳"这类空话。用这个结构：
"[身份标签]，[核心竞争力一句话]，[1-2个关键成果数据]。"

示例：
✅ "AI Agent开发实习生，精通LangChain+RAG技术栈，独立完成3个业务工具开发并将调用成功率从87%提升至96%。"
❌ "我是一名热爱编程的应届生，具有较强的学习能力和团队合作精神。"

=== 经历改写规则 ===
每条经历用STAR法则，**必须忠实于原文，只能润色不能编造**：
- company/role/duration：原封不动从原文提取，没有的留空，绝不编造
- situation: 一句话说清背景
- task: 一句话说清目标
- action: 用强动词开头的具体动作，基于原文已有的事实展开
- result: ⚠️ 必须有数字！原文有的保留，原文没有的根据上下文合理推估（如"提升近10%"→可写"10%"）

=== 技能归类规则 ===
- skills 字段分三层：核心技术栈 / 工具与平台 / 软技能
- 不要罗列十几个技能，列5-8个最有竞争力的
- 禁止使用"熟悉/了解/掌握"等程度词，只列技能名
- 原文没有的技能绝不添加，只能从经历描述中提取

=== 弱点诊断规则 ===
指出3个对HR最致命的硬伤，不要泛泛而谈。优先指出：
- 缺乏量化成果
- 关键词密度不足（过不了ATS筛选）
- 经历描述流水账、缺乏亮点
- 与目标岗位的技能缺口

输出纯JSON（不要markdown代码块、不要任何解释）：

{{
  "weaknesses": ["致命弱点1", "致命弱点2", "致命弱点3"],
  "suggestions": "有针对性的改进建议（2-3句话）",
  "optimized_resume": {{
    "name": "姓名",
    "job_target": "精确的求职意向（如：Python后端开发工程师、AI Agent开发工程师）",
    "contact": "联系方式（根据原文推断，格式：手机 | 邮箱 | 城市）",
    "education": "教育背景（完整格式：学校名 专业 学历·毕业时间）",
    "summary": "个人概述（2-3句，含身份标签+核心竞争力+关键数据）",
    "experiences": [
      {{
        "company": "公司/组织名",
        "role": "职位",
        "duration": "起止时间",
        "star_format": {{
          "situation": "一句话业务背景",
          "task": "一句话任务目标",
          "action": "2-3个强动词开头的具体动作，用分号分隔",
          "result": "可量化的成果（必须有数字！）"
        }}
      }}
    ],
    "skills": ["核心技术", "工具平台", "软技能"],
    "other": "加分项：证书（如CET-4/6分数）、比赛奖项、开源贡献、语言能力等"
  }}
}}

原始简历：
{resume_text}"""


def analyze_and_optimize(resume_text: str) -> dict:
    """一步完成简历分析+STAR优化"""
    print("=" * 50)
    print("[1/3] 分析简历并优化...")
    prompt = ANALYZE_PROMPT.replace("{resume_text}", resume_text)
    response = call_qwen(prompt)
    result = safe_json_parse(response, {"weaknesses": [], "suggestions": "", "optimized_resume": {}})
    print("  弱点:")
    for w in result.get("weaknesses", [])[:3]:
        print(f"    - {w}")
    print(f"  建议: {result.get('suggestions', '')}")
    return result


# ============ 第2步：本地模板 ============

def find_local_templates() -> list[Path]:
    """扫描 templates/ 下的所有模板文件（.docx / .doc）"""
    paths = []
    for ext in [".docx", ".doc"]:
        paths.extend(TEMPLATES_DIR.glob(f"*{ext}"))
    return sorted(paths)


# ============ 第3步：用户选择 ============

def user_select_template(paths: list[Path]) -> Path | None:
    """列出本地模板，等待用户选择"""
    print("\n" + "=" * 50)
    print("[2/3] 选择模板：")

    for i, p in enumerate(paths, 1):
        print(f"  [{i}] {p.name}")
    print(f"  [0] 不使用模板，AI从零生成")

    while True:
        try:
            choice = input("\n请输入编号: ").strip()
            num = int(choice)
            if num == 0:
                return None
            if 1 <= num <= len(paths):
                return paths[num - 1]
            print(f"  请输入0-{len(paths)}之间的数字")
        except ValueError:
            print("  请输入有效数字")


# ============ 第4步：AI逐段填入 ============

def _get_textboxes(doc: Document) -> list:
    """提取文档中所有文本框元素（wps:txbx / v:textbox），返回 (source_key, lxml_element) 列表"""
    import lxml.etree as ET
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'v': 'urn:schemas-microsoft-com:vml',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    body = doc.element.body
    textboxes = []

    # docx 文本框：wps:txbx / w:txbxContent
    for i, tb in enumerate(body.findall('.//w:txbxContent', nsmap)):
        textboxes.append({"index": None, "source": f"textbox_wps[{i}]", "element": tb, "text": _textbox_text(tb)})

    # 老版本 .doc 转过来的文本框：v:textbox
    for i, tb in enumerate(body.findall('.//v:textbox', nsmap)):
        textboxes.append({"index": None, "source": f"textbox_vml[{i}]", "element": tb, "text": _textbox_text(tb)})

    return textboxes


def _textbox_text(tb_element) -> str:
    """从文本框lxml元素提取纯文本"""
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = []
    for t in tb_element.iter(f'{{{w_ns}}}t'):
        if t.text:
            parts.append(t.text)
    return ''.join(parts)


def _replace_textbox_text(tb_element, new_text: str):
    """替换文本框内的文本，保留格式"""
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    runs = tb_element.findall(f'.//{{{w_ns}}}r')
    if runs:
        # 保留第一个 run 的格式，替换文本
        for i, r in enumerate(runs):
            if i == 0:
                t = r.find(f'{{{w_ns}}}t')
                if t is not None:
                    t.text = new_text
                    # 保留 rPr（格式属性）
                else:
                    from lxml.etree import SubElement
                    t = SubElement(r, f'{{{w_ns}}}t')
                    t.text = new_text
            else:
                r.getparent().remove(r)


def extract_paragraphs(doc: Document) -> list[dict]:
    """提取文档所有文本块：普通段落 + 表格 + 文本框"""
    items = []

    # 1. 正文段落
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            items.append({
                "index": len(items),
                "source": "paragraph",
                "text": para.text,
            })

    # 2. 表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    if para.text.strip():
                        items.append({
                            "index": len(items),
                            "source": f"table[{ti}].row[{ri}].cell[{ci}].p[{pi}]",
                            "text": para.text,
                        })

    # 3. 文本框
    for tb in _get_textboxes(doc):
        text = tb["text"].strip()
        if text:
            items.append({
                "index": len(items),
                "source": tb["source"],
                "text": text,
            })

    return items


def apply_replacements(doc: Document, replacements: list[dict]):
    """按映射原位替换文本，保留原有格式"""
    # 构建索引映射：paragraph items + table items + textbox items
    para_items = []
    table_items = {}
    textbox_items = {}

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            para_items.append({"para": para, "index": len(para_items)})

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    if para.text.strip():
                        table_items[f"table[{ti}].row[{ri}].cell[{ci}].p[{pi}]"] = para

    for tb in _get_textboxes(doc):
        textbox_items[tb["source"]] = tb["element"]

    for rep in replacements:
        action = rep.get("action", "keep")
        if action != "replace":
            continue

        new_text = rep.get("new_text", "")
        idx = rep.get("index", -1)
        source = rep.get("source", "")

        if source.startswith("textbox_"):
            tb_elem = textbox_items.get(source)
            if tb_elem is not None:
                _replace_textbox_text(tb_elem, str(new_text))
        elif "table" in source:
            target = table_items.get(source)
            if target:
                _replace_para(target, str(new_text))
        else:
            target = para_items[idx]["para"] if idx < len(para_items) else None
            if target:
                _replace_para(target, str(new_text))


def _replace_para(para, new_text: str):
    """替换段落文本，保留第一个run的格式"""
    if para.runs:
        font = para.runs[0].font
        para.clear()
        new_run = para.add_run(new_text)
        new_run.font.name = font.name
        new_run.font.size = font.size
        new_run.font.bold = font.bold
        new_run.font.italic = font.italic
        new_run.font.color.rgb = font.color.rgb
    else:
        para.text = new_text


def generate_from_scratch(data: dict) -> Path:
    """降级方案：AI从零生成一个简洁专业的简历.docx"""
    # ponytail: 用AI生成docx结构，再逐段写入，不是渲染HTML
    # 直接用python-docx手动构建一份专业简历

    doc = Document()
    opt = data.get("optimized_resume", data)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题：姓名
    name = opt.get("name", "姓名")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(name)
    run.font.size = Pt(22)
    run.font.bold = True

    # 求职意向 + 联系方式一行
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = opt.get("job_target", "")
    if opt.get("education"):
        edu = opt["education"]
        info_text += f"  |  {edu}" if info_text else str(edu)
    run = info_para.add_run(info_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # 分割线
    doc.add_paragraph("─" * 50)

    # 个人概述
    summary = opt.get("summary", "")
    if summary:
        h = doc.add_paragraph()
        run = h.add_run("个人概述")
        run.font.size = Pt(14)
        run.font.bold = True
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(6)

    # 经历
    exps = opt.get("experiences", [])
    if exps:
        h = doc.add_paragraph()
        run = h.add_run("项目经历")
        run.font.size = Pt(14)
        run.font.bold = True
        for exp in exps:
            # 公司 + 职位 + 时间
            header = f"{exp.get('role', '')}  |  {exp.get('company', '')}  |  {exp.get('duration', '')}"
            hp = doc.add_paragraph()
            r = hp.add_run(header)
            r.font.size = Pt(11)
            r.font.bold = True

            # STAR内容
            sf = exp.get("star_format", {})
            for label, key in [("S", "situation"), ("T", "task"), ("A", "action"), ("R", "result")]:
                val = sf.get(key, "")
                if val:
                    bp = doc.add_paragraph()
                    bp.paragraph_format.left_indent = Cm(0.5)
                    r = bp.add_run(f"{label}: {val}")
                    r.font.size = Pt(10)

    # 技能
    skills = opt.get("skills", [])
    if skills:
        h = doc.add_paragraph()
        run = h.add_run("专业技能")
        run.font.size = Pt(14)
        run.font.bold = True
        sp = doc.add_paragraph()
        r = sp.add_run(" · ".join(skills))
        r.font.size = Pt(10)

    # 其他
    other = opt.get("other", "")
    if other:
        h = doc.add_paragraph()
        run = h.add_run("其他")
        run.font.size = Pt(14)
        run.font.bold = True
        op = doc.add_paragraph(other)
        op.paragraph_format.space_after = Pt(6)

    path = OUTPUT_DIR / "resume_from_scratch.docx"
    doc.save(str(path))
    return path


def convert_doc_to_docx(doc_path: Path) -> Path | None:
    """将 .doc 转为 .docx：优先用 Word COM，其次 LibreOffice，否则提示用户手动转换"""
    import subprocess

    # 方案1：Windows Word COM 自动化
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(doc_path.resolve()))
        converted = TEMPLATES_DIR / f"{doc_path.stem}.docx"
        doc.SaveAs2(str(converted.resolve()), FileFormat=16)  # 16 = wdFormatXMLDocument
        doc.Close()
        word.Quit()
        if converted.exists():
            print(f"  .doc 已自动转换为 .docx")
            return converted
    except Exception:
        pass

    # 方案2：LibreOffice headless
    import shutil
    lo_paths = [
        "soffice",
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
    ]
    for lo in lo_paths:
        if shutil.which(lo) or Path(lo).exists():
            result = subprocess.run(
                [lo, "--headless", "--convert-to", "docx", str(doc_path), "--outdir", str(TEMPLATES_DIR)],
                capture_output=True, text=True, timeout=30,
            )
            converted = TEMPLATES_DIR / f"{doc_path.stem}.docx"
            if converted.exists():
                print(f"  .doc 已转换为 .docx")
                return converted
            break

    # 两种方案都不可用
    print(f"  警告：无法自动转换 .doc 文件")
    print(f"  请用 Word/WPS 打开 {doc_path.name}，另存为 .docx 格式放入 templates/")
    return None


def open_document(path: Path):
    """打开 .docx（必要时自动转换 .doc），返回 Document 或 None"""
    if path.suffix.lower() == ".doc":
        converted = convert_doc_to_docx(path)
        return Document(str(converted)) if converted else None
    try:
        return Document(str(path))
    except Exception as e:
        print(f"  打开模板失败: {e}")
        return None


def fill_template(template_path: Path, optimized_data: dict) -> Path:
    """AI逐段分析模板并填入内容"""
    print("\n" + "=" * 50)
    print("[3/3] AI分析模板并填入内容...")

    doc = open_document(template_path)
    if doc is None:
        return generate_from_scratch(optimized_data)
    items = extract_paragraphs(doc)

    if not items:
        print("  模板无有效文本，fallback到从零生成")
        return generate_from_scratch(optimized_data)

    # 构造AI prompt
    items_text = "\n".join(
        f'{{"index": {it["index"]}, "source": "{it["source"]}", "text": "{it["text"]}"}}'
        for it in items
    )
    resume_json = json.dumps(optimized_data.get("optimized_resume", optimized_data), ensure_ascii=False, indent=2)

    fill_prompt = f"""你是简历排版专家，为HR优化过1000+份简历。下面是一个Word模板的所有文本块，和一份优化后的简历。

请逐块判断替换策略，输出JSON数组。

=== 替换规则 ===
1. 示例姓名→真实姓名。模板中的"张三/李四/王五"等示例名一律替换
2. 示例学校→真实学校+专业+学历
3. 示例经历→优化后的经历，每段经历拆成要点分号分隔，每条≤30字
4. 示例技能→优化后的技能，用"·"连接，不超过8个
5. summary写入个人概述，简洁有力
6. 章节标题（教育背景/工作经历/技能特长/项目经验/个人概述等）→保留不动
7. 模板装饰文字（"自我评价""联系方式"等栏目标签）→保留不动
8. 与简历内容无关的模板说明/引导语→保留不动

=== 撰写风格 ===
- 每个文本块输入内容精简到15-30字，适配模板空间
- 经历描述用"动词+对象+成果（数字）"格式
- 禁止空洞形容词（丰富的、较强的、良好的），用数字说话

模板文本块：
{items_text}

优化后的简历：
{resume_json}

只输出JSON数组，不要任何解释、不要markdown代码块：
[{{"index": 0, "source": "...", "action": "replace|keep", "new_text": "..."}}]"""

    try:
        response = call_qwen(fill_prompt)
        replacements = safe_json_parse(response, [])
        if not replacements:
            print("  AI未返回有效映射，fallback到从零生成")
            return generate_from_scratch(optimized_data)

        print(f"  共 {len([r for r in replacements if r.get('action') == 'replace'])} 处替换")

        apply_replacements(doc, replacements)

        path = OUTPUT_DIR / "resume_final.docx"
        doc.save(str(path))
        return path

    except Exception as e:
        print(f"  填入失败: {e}，fallback到从零生成")
        return generate_from_scratch(optimized_data)


def extract_pdf_text(pdf_path: Path) -> str:
    """从PDF提取纯文本"""
    doc = fitz.open(str(pdf_path))
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return text.strip()


# ============ 主流程 ============

def main():
    # 读取简历：扫描 resumes/ 下的 .pdf 或 .txt
    for ext in [".pdf", ".txt"]:
        candidates = list(RESUMES_DIR.glob(f"*{ext}"))
        if candidates:
            resume_path = candidates[0]
            break
    else:
        print(f"请将简历文件（.pdf 或 .txt）放入 {RESUMES_DIR}")
        sys.exit(1)

    if resume_path.suffix.lower() == ".pdf":
        print(f"读取PDF简历: {resume_path}")
        resume_text = extract_pdf_text(resume_path)
        if not resume_text:
            print("PDF中未提取到文字，请确认文件不是扫描件")
            sys.exit(1)
        print(f"提取到 {len(resume_text)} 个字")
    else:
        resume_text = resume_path.read_text(encoding="utf-8")
        print(f"读取简历: {resume_path}")

    # 第1步：分析 + 优化
    result = analyze_and_optimize(resume_text)
    if "error" in result:
        print(f"分析失败: {result}")
        sys.exit(1)

    # 第2步：本地模板
    print("\n" + "=" * 50)
    print("[2/3] 加载本地模板...")
    templates = find_local_templates()
    if not templates:
        print("  模板文件夹为空，直接AI从零生成")
        print(f"  （可将.docx模板放入 {TEMPLATES_DIR} 目录）")
        path = generate_from_scratch(result)
        print(f"\n[完成] 简历已生成: {path}")
        return

    # 第3步：用户选择
    chosen = user_select_template(templates)
    if chosen is None:
        print("  跳过模板，AI从零生成")
        path = generate_from_scratch(result)
        print(f"\n[完成] 简历已生成: {path}")
        return

    # AI填入
    path = fill_template(chosen, result)
    print(f"\n[完成] 简历已生成: {path}")


if __name__ == "__main__":
    main()
